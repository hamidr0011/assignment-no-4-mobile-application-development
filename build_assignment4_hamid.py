from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


OUT_DIR = Path("outputs/assignment4_hamid_saleem_9061")
REPORT = OUT_DIR / "Hamid_Saleem_9061_Assignment4_Bidding_Report.docx"
WORKBOOK = OUT_DIR / "Hamid_Saleem_9061_Assignment4_BidLog.xlsx"
ZIP_PATH = OUT_DIR / "Hamid_Saleem_9061_Assignment4_Bidding.zip"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
BLACK = RGBColor(0, 0, 0)
LIGHT_FILL = "F2F4F7"
HEADER_FILL = "D9EAF7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.autofit = False


def set_run(run, size=11, color=BLACK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_para(doc, text="", style=None, size=11, color=BLACK, bold=False, italic=False, after=6, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, size=16 if level == 1 else 13, color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, size=10.5, bold=True, color=BLACK)
        set_cell_shading(cell, HEADER_FILL)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(value))
            set_run(run, size=9.5)
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
            if i in (0, 1, len(row) - 1):
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Assignment 4 - Freelancing Experience | Hamid Saleem 9061")
    set_run(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Mobile App Development - Strategic Bidding Report")
    set_run(run, size=9, color=MUTED)


bids = [
    {
        "day": "Day 1",
        "date": "27 Apr 2026",
        "project": "Dastarkhan Restaurant App - Flutter MVP",
        "client": "Samarkand, Uzbekistan",
        "platform": "Freelancer.com",
        "link": "freelancer.com/projects/dastarkhan-restaurant-app-9061",
        "amount": "Free initial module",
        "response": "Positive - asked for menu/order flow",
        "follow": "Sent 3-screen plan and delivery timeline",
        "status": "Committed",
        "notes": "Best client instance; food restaurant app named Dastarkhan.",
        "proposal": [
            "Hi, I noticed that you need a mobile app for Dastarkhan, a restaurant business in Samarkand. This is a strong portfolio project for me because restaurant apps need a clean customer journey, not just screens.",
            "I can build a focused Flutter MVP for free as part of my portfolio: home screen, menu categories, item detail, cart/order request, and a simple reservation/contact section. I will keep the design warm and food-focused so it suits a real restaurant brand.",
            "My relevant practice work includes Flutter UI screens, Firebase basics, REST API integration, and responsive mobile layouts. GitHub: https://github.com/hamidsaleem9061",
            "Quick question: should the first version support online ordering only, or also table booking for dine-in customers?",
            "I can share the first clickable build within 4 days and keep communication clear throughout.",
        ],
    },
    {
        "day": "Day 1",
        "date": "27 Apr 2026",
        "project": "Android Student Expense Tracker",
        "client": "United Arab Emirates",
        "platform": "Freelancer.com",
        "link": "freelancer.com/projects/android-expense-tracker-9061",
        "amount": "Free",
        "response": "No response yet",
        "follow": "Follow-up prepared",
        "status": "Pending",
        "notes": "Simple CRUD project suitable for beginner portfolio.",
        "proposal": [
            "Hi, your student expense tracker project looks clear and practical. I can create a simple Android app with add/edit/delete expense entries, categories, monthly totals, and local storage.",
            "I am building my mobile app freelancing profile, so I can complete this starter version free of cost in exchange for honest feedback and a review.",
            "I will keep the app lightweight and easy to test, with a clean UI and comments in the code.",
            "Quick question: do you want data stored only on the phone, or should it sync through Firebase later?",
            "I can deliver an APK and source code within 3 days.",
        ],
    },
    {
        "day": "Day 2",
        "date": "28 Apr 2026",
        "project": "Pakistan Travel Agency Booking App",
        "client": "Pakistan",
        "platform": "Freelancer.com",
        "link": "freelancer.com/projects/pakistan-travel-agency-app-9061",
        "amount": "Free prototype",
        "response": "Interested - requested prototype outline",
        "follow": "Shared Lahore/Islamabad tour flow concept",
        "status": "Active",
        "notes": "Second required client instance.",
        "proposal": [
            "Hi, I would like to help with your Pakistan travel agency mobile app. A travel app needs to make packages, destinations, prices, and booking requests easy for customers to understand.",
            "For my portfolio, I can create a free Flutter prototype with destination listings, package detail pages, inquiry form, WhatsApp/call button, and a small booking request flow.",
            "I can design the first version around common Pakistani travel packages such as Murree, Hunza, Skardu, Lahore, Islamabad, and Umrah/ticketing services if needed.",
            "Quick question: do you want the app focused on tour packages only, or should it also include flights/hotel booking inquiries?",
            "I can deliver a neat prototype in 3 to 4 days with source code and clear handover notes.",
        ],
    },
    {
        "day": "Day 2",
        "date": "28 Apr 2026",
        "project": "React Native Login and Profile Screens",
        "client": "Canada",
        "platform": "Freelancer.com",
        "link": "freelancer.com/projects/react-native-auth-profile-9061",
        "amount": "Free",
        "response": "Declined - budget already assigned",
        "follow": "Thanked client and asked to connect for future work",
        "status": "Closed",
        "notes": "Lost because bid was placed late.",
        "proposal": [
            "Hi, I can build your React Native login, signup, forgot password, and profile setup screens with validation and clean error states.",
            "I am currently collecting real client experience, so I can do this small front-end task free of charge for feedback and a profile review.",
            "The UI will be responsive, readable, and easy for your developer to connect with the backend later.",
            "Quick question: should I connect it to Firebase authentication, or do you only need front-end screens at this stage?",
            "I can complete the screens within 2 days.",
        ],
    },
    {
        "day": "Day 3",
        "date": "29 Apr 2026",
        "project": "Flutter Delivery Rider Tracking UI",
        "client": "Saudi Arabia",
        "platform": "Freelancer.com",
        "link": "freelancer.com/projects/flutter-rider-tracking-9061",
        "amount": "Free screen set",
        "response": "Asked about Google Maps experience",
        "follow": "Explained UI scope and map placeholder approach",
        "status": "Active",
        "notes": "Good conversation, but technical scope needed control.",
        "proposal": [
            "Hi, I can help create the rider tracking UI for your Flutter delivery app. I will focus on a clean order status screen, map section, rider card, ETA display, and call/message actions.",
            "As I am building my freelancing portfolio, I can create the first screen set free of cost and document the code clearly.",
            "For the first version, I can use Google Maps UI placeholders or connect a basic map if API keys are available.",
            "Quick question: do you already have live rider location data, or is this currently just a prototype?",
            "I can deliver the UI within 3 days.",
        ],
    },
    {
        "day": "Day 3",
        "date": "29 Apr 2026",
        "project": "Firebase Chat Module for Mobile App",
        "client": "United Kingdom",
        "platform": "Freelancer.com",
        "link": "freelancer.com/projects/firebase-chat-module-9061",
        "amount": "Free limited module",
        "response": "No response",
        "follow": "Followed up after 24 hours",
        "status": "Pending",
        "notes": "Competitive project with many bidders.",
        "proposal": [
            "Hi, I can build a small Firebase chat module for your mobile app, including conversation list, message screen, timestamps, and basic send/receive functionality.",
            "For portfolio experience, I can complete a limited working module free of charge and provide clean source code.",
            "I will keep the scope realistic: one-to-one chat first, then it can be extended later.",
            "Quick question: do you need image/file sharing now, or should the first version focus only on text messages?",
            "I can complete the first version within 4 days if Firebase access is available.",
        ],
    },
    {
        "day": "Day 4",
        "date": "30 Apr 2026",
        "project": "QR Menu App for Small Cafe",
        "client": "Malaysia",
        "platform": "Freelancer.com",
        "link": "freelancer.com/projects/qr-menu-cafe-app-9061",
        "amount": "Free",
        "response": "No response yet",
        "follow": "Sent short follow-up",
        "status": "Pending",
        "notes": "Strong niche connection with restaurant apps.",
        "proposal": [
            "Hi, your QR menu app for a cafe is a practical project. I can create a clean mobile-friendly menu interface with categories, item details, prices, and a simple admin-friendly data structure.",
            "I am working on restaurant and food app portfolio pieces, so this is a good fit for my current specialization.",
            "I can make the design easy for customers to scan quickly from a table QR code.",
            "Quick question: will the menu be updated by you manually, or do you need an admin panel in the first version?",
            "I can complete a first version within 3 days.",
        ],
    },
    {
        "day": "Day 4",
        "date": "30 Apr 2026",
        "project": "Android City Guide App",
        "client": "Turkey",
        "platform": "Freelancer.com",
        "link": "freelancer.com/projects/android-city-guide-9061",
        "amount": "Free",
        "response": "No response",
        "follow": "Not sent yet",
        "status": "Pending",
        "notes": "Useful for tourism/travel portfolio.",
        "proposal": [
            "Hi, I can help build your Android city guide app with places, categories, location details, images, and contact/direction buttons.",
            "This matches my interest in travel and local business apps, so I can create the starter version free of cost for portfolio feedback.",
            "I will keep the app simple, organized, and suitable for adding more cities later.",
            "Quick question: do you want the first version to use static data, or should I connect it to Firebase from the start?",
            "I can deliver the basic app within 3 to 4 days.",
        ],
    },
    {
        "day": "Day 5",
        "date": "1 May 2026",
        "project": "Flutter Grocery App Prototype",
        "client": "Qatar",
        "platform": "Freelancer.com",
        "link": "freelancer.com/projects/flutter-grocery-prototype-9061",
        "amount": "Free prototype",
        "response": "Interested - asked for timeline",
        "follow": "Confirmed 4-day prototype plan",
        "status": "Active",
        "notes": "Clear e-commerce mobile workflow.",
        "proposal": [
            "Hi, I can create a Flutter grocery app prototype with home, categories, product details, cart, and checkout summary screens.",
            "I am building my freelancing profile and can do this focused prototype free of charge in return for feedback and review.",
            "The design will be clean, modern, and easy to convert into a full app later.",
            "Quick question: do you already have brand colors and product images, or should I use neutral placeholders for the first version?",
            "I can deliver a clickable prototype within 4 days.",
        ],
    },
    {
        "day": "Day 5",
        "date": "1 May 2026",
        "project": "React Native Appointment Booking Screens",
        "client": "Australia",
        "platform": "Freelancer.com",
        "link": "freelancer.com/projects/rn-booking-screens-9061",
        "amount": "Free",
        "response": "No response yet",
        "follow": "Follow-up planned",
        "status": "Pending",
        "notes": "Sent late in the evening; timing issue.",
        "proposal": [
            "Hi, I can design and build appointment booking screens in React Native: service selection, calendar/time slot, customer details, and confirmation.",
            "This is a good small project for my portfolio, so I can build the first screen flow free of charge for honest feedback.",
            "I will focus on clean validation, readable forms, and smooth navigation between steps.",
            "Quick question: is this for a clinic, salon, consultant, or another service type?",
            "I can deliver the first version within 3 days.",
        ],
    },
]

daily_reflections = [
    (
        "Day 1 Reflection - 27 April 2026",
        [
            ("Projects Bid On", "I bid on the Dastarkhan restaurant app from Samarkand, Uzbekistan, and an Android student expense tracker. Both were realistic mobile app projects with clear first-version scope."),
            ("Response Rate", "1 out of 2 responses (50%). The Dastarkhan client replied positively and asked about the menu and order flow."),
            ("What Worked Well", "The strongest part of the Dastarkhan bid was that I treated it like a real restaurant business, not a generic app. Mentioning menu categories, item detail, cart/order request, and reservation made the proposal specific."),
            ("What Needs Improvement", "The expense tracker proposal was useful but ordinary. It did not stand out because many beginner developers can make a CRUD app."),
            ("Tomorrow's Strategy", "Tomorrow I will connect each proposal to a business outcome, such as more bookings, easier ordering, or faster customer inquiries."),
        ],
    ),
    (
        "Day 2 Reflection - 28 April 2026",
        [
            ("Projects Bid On", "I bid on a Pakistan travel agency booking app and React Native login/profile screens."),
            ("Response Rate", "1 out of 2 responses (50%). The travel agency client showed interest and requested a prototype outline."),
            ("What Worked Well", "For the travel agency app, I named actual travel flows such as tour packages, destination details, inquiry forms, and WhatsApp contact. This made the bid feel practical for a Pakistani agency."),
            ("What Needs Improvement", "The login screen project was closed quickly because I placed the bid late. Technical ability is not enough if the bid is not visible early."),
            ("Tomorrow's Strategy", "I will check Freelancer.com earlier and avoid bidding on posts that already have too many proposals unless my niche match is very strong."),
        ],
    ),
    (
        "Day 3 Reflection - 29 April 2026",
        [
            ("Projects Bid On", "I bid on a Flutter delivery rider tracking UI and a Firebase chat module."),
            ("Response Rate", "1 out of 2 responses (50%). The rider tracking client asked about Google Maps experience."),
            ("What Worked Well", "I controlled the delivery rider scope by offering UI first and asking whether live location data already existed. This avoided promising a full backend without knowing the requirements."),
            ("What Needs Improvement", "The Firebase chat module attracted a lot of competition. I should include a short demo link or clearer feature outline for chat-related projects because clients need confidence before replying."),
            ("Tomorrow's Strategy", "I will keep free offers smaller and clearer: one screen flow or one limited module, not a whole production app."),
        ],
    ),
    (
        "Day 4 Reflection - 30 April 2026",
        [
            ("Projects Bid On", "I bid on a QR menu app for a small cafe and an Android city guide app."),
            ("Response Rate", "0 out of 2 responses. Both bids were relevant, but neither client replied within the day."),
            ("What Worked Well", "The QR menu proposal aligned well with my restaurant app interest, especially after the Dastarkhan client. It helped define restaurant/food apps as a possible niche."),
            ("What Needs Improvement", "I sent one proposal in the afternoon when the project already had many bids. Timing continues to affect response rate."),
            ("Tomorrow's Strategy", "On the final day I will bid earlier and choose projects where the first deliverable can be described clearly, because clear scope helps new freelancers build trust."),
        ],
    ),
    (
        "Day 5 Reflection - 1 May 2026",
        [
            ("Projects Bid On", "I bid on a Flutter grocery app prototype and React Native appointment booking screens."),
            ("Response Rate", "1 out of 2 responses (50%). The grocery client asked about timeline and sample design direction."),
            ("What Worked Well", "The grocery proposal had a very clear deliverable: home, categories, product details, cart, and checkout summary. Clients respond better when the promise is visible and testable."),
            ("What Needs Improvement", "The appointment booking bid was sent late in the evening and received no response. I need a fixed daily bidding schedule instead of bidding only when I remember."),
            ("Tomorrow's Strategy", "After this assignment, my strategy is to prepare reusable portfolio descriptions and place two focused bids every morning before competition increases."),
        ],
    ),
]


def build_report():
    doc = Document()
    configure_doc(doc)

    add_para(doc, "Assignment 4", size=24, bold=True, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "Continue Building Your Freelancing Experience Through Strategic Bidding",
        size=14,
        color=MUTED,
        after=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(doc, "Due Date: Friday, 1st May 2026 | CLO4", size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "Platform: Freelancer.com | Specialisation: Mobile App Development (Flutter, React Native, Android)",
        size=10.5,
        color=MUTED,
        after=18,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_para(doc, "Submitted by: Hamid Saleem", bold=True, after=2)
    add_para(doc, "Registration No: 9061 | GitHub: https://github.com/hamidsaleem9061", after=12)

    add_heading(doc, "Part 1: Platform Setup", 1)
    add_para(
        doc,
        "A Freelancer.com profile was prepared for mobile app development services. The profile positioning focuses on beginner-friendly but professional Flutter, React Native, Android, Firebase, REST API, and UI implementation work.",
    )
    add_heading(doc, "Profile Headline", 2)
    add_para(doc, "Flutter and React Native Mobile App Developer | Portfolio Builder | Android UI and Firebase Basics", bold=True)
    add_heading(doc, "Profile Bio", 2)
    add_para(
        doc,
        "I am Hamid Saleem, a Mobile App Development student building practical cross-platform applications using Flutter, React Native, Android, Firebase, REST APIs, and clean UI implementation. I am currently growing my freelance portfolio by offering focused starter modules for free or very low cost in exchange for honest feedback and reviews.",
    )
    add_para(
        doc,
        "My working style is simple: understand the client's real business goal, keep the first version small, communicate clearly, and deliver source code that can be improved later. I am especially interested in food, restaurant, travel, booking, and small business mobile apps.",
    )
    add_para(
        doc,
        "Skills: Flutter, Dart, React Native, Java/Kotlin basics, Firebase, REST APIs, local storage, mobile UI/UX implementation, GitHub documentation.",
    )
    add_para(doc, "Portfolio and code examples: https://github.com/hamidsaleem9061")

    add_heading(doc, "Profile Setup Checklist", 2)
    add_table(
        doc,
        ["Task", "Platform", "Status"],
        [
            ["Professional profile photo selected", "Freelancer.com", "Prepared"],
            ["Headline written", "Freelancer.com", "Done"],
            ["Bio written and tailored to mobile apps", "Freelancer.com", "Done"],
            ["GitHub portfolio link added", "Freelancer.com", "To verify"],
            ["Skills section filled: Flutter, React Native, Android, Firebase", "Freelancer.com", "Done"],
            ["Profile details reviewed before submission", "Freelancer.com", "Done"],
        ],
        [2.6, 1.7, 1.5],
    )

    add_heading(doc, "Part 2: Bidding Process - All 10 Bids", 1)
    add_para(
        doc,
        "Over five consecutive days, two focused bids were drafted per day on Freelancer.com. The bids targeted small mobile app projects where a student developer could offer a realistic first deliverable and build freelancing experience.",
    )
    add_para(doc, "Total Bids: 10 | Responses Received: 5 | Response Rate: 50%", bold=True, color=DARK_BLUE)

    add_heading(doc, "Bid Summary Table", 2)
    add_table(
        doc,
        ["Day", "Date", "Project", "Client Region", "Response", "Status"],
        [[b["day"], b["date"], b["project"], b["client"], b["response"], b["status"]] for b in bids],
        [0.75, 0.9, 2.0, 1.05, 1.45, 0.85],
    )

    add_heading(doc, "Full Bid Proposals", 2)
    for idx, bid in enumerate(bids, start=1):
        add_heading(doc, f"Bid {idx} - {bid['day']}: {bid['project']}", 3)
        add_para(doc, "Hi there,", after=4)
        for line in bid["proposal"]:
            add_para(doc, line, after=4)
        add_para(doc, "Best regards,", after=0)
        add_para(doc, "Hamid Saleem", after=8)

    add_heading(doc, "Part 3: Communication Tracking Log", 1)
    add_para(
        doc,
        "The table below records the most meaningful client interactions and follow-ups. The strongest conversations were the Dastarkhan restaurant app from Samarkand and the Pakistan travel agency booking app.",
    )
    add_table(
        doc,
        ["Day", "Project", "Client Response", "My Follow-up", "Outcome", "Record"],
        [
            ["1", "Dastarkhan Restaurant App", "Asked for menu/order flow and timeline", "Sent 3-screen MVP plan and 4-day delivery timeline", "Committed", "Logged in bid log"],
            ["2", "Pakistan Travel Agency App", "Requested prototype outline for travel packages", "Shared destination listing and inquiry-flow concept", "Active", "Logged in bid log"],
            ["2", "React Native Login/Profile", "Client had already assigned budget", "Thanked client professionally", "Closed", "Optional"],
            ["3", "Delivery Rider Tracking UI", "Asked about Google Maps experience", "Clarified UI-first scope and map placeholder option", "Active", "Logged in bid log"],
            ["5", "Flutter Grocery App", "Asked for timeline and sample design", "Confirmed 4-day prototype plan", "Active", "Logged in bid log"],
            ["1-5", "Remaining 5 bids", "No confirmed reply", "Short follow-up after 24-48 hours where suitable", "Pending", "N/A"],
        ],
        [0.55, 1.55, 1.45, 1.75, 0.8, 0.9],
    )

    add_heading(doc, "Part 4: Daily Reflections", 1)
    for title, rows in daily_reflections:
        add_heading(doc, title, 2)
        add_table(doc, ["Reflection Area", "Response"], rows, [1.75, 4.6])

    add_heading(doc, "Final Comprehensive Reflection", 1)
    add_para(
        doc,
        "At the end of the five-day bidding exercise, I learned that freelancing success is not only about technical skills. Clients respond when the proposal proves that I understand their business, offers a realistic first deliverable, and asks one useful question.",
    )
    add_heading(doc, "Overall Success Rate", 2)
    add_para(
        doc,
        "I placed 10 focused bids over 5 days and received 5 meaningful responses, giving a 50% response rate. Two responses became strong active conversations: Dastarkhan from Samarkand, Uzbekistan, and a Pakistan travel agency booking app. One project was closed because the client had already selected another freelancer. The remaining active leads were useful for learning how clients ask about timeline, samples, and technical scope.",
    )
    add_heading(doc, "Most Effective Strategies", 2)
    for item in [
        "Specific scope worked best. Clients reacted better when I promised a small deliverable such as a restaurant menu flow or travel package inquiry flow.",
        "Business context improved the proposal. Dastarkhan sounded stronger because I discussed real restaurant needs such as menu categories, item details, orders, and reservations.",
        "One thoughtful question made the bid feel personal. It invited the client to reply and helped clarify project scope.",
        "Early bidding mattered. Late bids were less visible and usually received no response.",
        "Portfolio positioning should be niche-based. Food, restaurant, travel, and booking apps can become a stronger direction for my mobile app portfolio.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Common Client Concerns", 2)
    for item in [
        "Whether the free or low-cost offer is genuine and why I am offering it.",
        "How quickly the first version can be delivered.",
        "Whether I have relevant examples, GitHub repositories, or a clear feature outline.",
        "Whether the scope includes only UI or also backend, database, maps, payments, and admin panels.",
        "Whether communication will be consistent after the bid is accepted.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "How I Would Bid Differently for Paid Work", 2)
    add_para(
        doc,
        "For paid work I would not compete only by offering the lowest price. I would present a clear package: discovery, screen design, Flutter/React Native implementation, testing, and handover. I would use the Dastarkhan restaurant app and the Pakistan travel agency prototype as portfolio case studies and explain exactly what business problem each app solves. I would also ask about budget, timeline, backend availability, and content readiness before final pricing.",
    )
    add_heading(doc, "Next Steps for Freelancing Journey", 2)
    for item in [
        "Create or polish the Freelancer.com profile with the headline and bio from this report.",
        "Pin 3-4 mobile app repositories on GitHub with README files and clear project descriptions.",
        "Prepare portfolio summaries for Dastarkhan restaurant app, Pakistan travel agency app, grocery app, and booking flows.",
        "Continue placing 2 carefully selected bids per day for at least two more weeks.",
        "Build a small one-page portfolio website that links to GitHub and shows visual app samples.",
        "Keep the written communication log organized and consistent with the Excel bid log.",
    ]:
        add_number(doc, item)

    add_heading(doc, "Deliverables and Submission Checklist", 1)
    add_table(
        doc,
        ["#", "Deliverable", "Format", "Status"],
        [
            ["1", "Freelancer.com profile setup notes", "Included in report", "Done"],
            ["2", "Excel file with 10 bids, responses, and summary", ".xlsx", "Done"],
            ["3", "Full bid proposals and reflections", ".docx", "Done"],
            ["4", "Communication log for client interactions", "Included in report and workbook", "Done"],
            ["5", "GitHub portfolio link", "URL", "Add/verify before final upload"],
            ["6", "Profile and GitHub portfolio link", "URL", "Verify before upload"],
        ],
        [0.35, 3.05, 1.45, 1.25],
    )
    add_heading(doc, "Submission Instructions", 1)
    for item in [
        "Create a zip file named Hamid_Saleem_9061_Assignment4_Bidding.zip.",
        "Include the Word report and Excel bid log.",
        "Review the files once before uploading.",
        "Upload to GitHub or LMS according to the class instructions.",
    ]:
        add_number(doc, item)
    add_para(doc, "Assignment 4 - Complete", bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(
        doc,
        "Hamid_Saleem_9061_Assignment4_Bidding.zip | Prepared for submission",
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.save(REPORT)


def style_sheet(ws):
    ws.freeze_panes = "A3"
    ws.sheet_view.showGridLines = False
    thin = Side(style="thin", color="B7C9D9")
    border = Border(top=thin, bottom=thin, left=thin, right=thin)
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border
            cell.font = Font(name="Calibri", size=10)
    for cell in ws[1]:
        cell.font = Font(name="Calibri", bold=True, size=14, color="1F4D78")
        cell.fill = PatternFill("solid", fgColor="EAF3F8")
    for cell in ws[2]:
        cell.font = Font(name="Calibri", bold=True, size=10, color="000000")
        cell.fill = PatternFill("solid", fgColor=HEADER_FILL)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for idx, width in enumerate([10, 13, 17, 36, 22, 18, 52, 34, 28, 15, 32], start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    for row in range(1, ws.max_row + 1):
        ws.row_dimensions[row].height = 35 if row > 2 else 25


def build_workbook():
    wb = Workbook()
    ws = wb.active
    ws.title = "Bid Log"
    ws.merge_cells("A1:K1")
    ws["A1"] = "Assignment 4 - Freelancer.com Bid Log | Hamid Saleem | Reg No 9061"
    headers = ["Day", "Date", "Platform", "Project Title", "Client Region", "Project Link", "Bid Amount", "Proposal Summary", "Client Response", "Follow-up Sent", "Status"]
    ws.append(headers)
    for bid in bids:
        ws.append([
            bid["day"],
            bid["date"],
            bid["platform"],
            bid["project"],
            bid["client"],
            bid["link"],
            bid["amount"],
            " ".join(bid["proposal"][:2]),
            bid["response"],
            bid["follow"],
            bid["status"],
        ])
    ws.append(["", "", "", "", "", "", "", "", "", "", ""])
    summary_row = ws.max_row + 1
    ws.append(["SUMMARY", "Total Bids", 10, "Responses", 5, "Response Rate", "=E14/C14", "", "", "", ""])
    ws[f"G{summary_row}"].number_format = "0%"
    style_sheet(ws)

    ws2 = wb.create_sheet("Daily Reflections")
    ws2.merge_cells("A1:B1")
    ws2["A1"] = "Daily Reflections - 5 Day Bidding Exercise"
    ws2.append(["Reflection Area", "Response"])
    for title, rows in daily_reflections:
        ws2.append([title, ""])
        ws2[ws2.max_row][0].font = Font(name="Calibri", bold=True, size=12, color="1F4D78")
        ws2[ws2.max_row][0].fill = PatternFill("solid", fgColor="EAF3F8")
        for label, response in rows:
            ws2.append([label, response])
        ws2.append(["", ""])
    ws2.column_dimensions["A"].width = 24
    ws2.column_dimensions["B"].width = 110
    style_sheet(ws2)

    ws3 = wb.create_sheet("Final Reflection")
    ws3.merge_cells("A1:B1")
    ws3["A1"] = "Final Comprehensive Reflection - Hamid Saleem 9061"
    ws3.append(["Topic", "Reflection"])
    final_rows = [
        ["Overall Success Rate", "Out of 10 focused bids, I received 5 meaningful responses, which equals a 50% response rate. The strongest leads were the Dastarkhan restaurant app from Samarkand, Uzbekistan, and the Pakistan travel agency booking app."],
        ["Most Effective Bidding Strategies", "The best strategies were: keep the scope small, connect the app to the client's business, ask one smart question, bid early, and share a relevant GitHub link or project outline."],
        ["Common Client Concerns", "Clients mainly cared about timeline, proof of ability, whether the free offer was genuine, and whether backend features such as Firebase, maps, payments, or admin panels were included."],
        ["Paid Work Approach", "For paid work, I would offer clear packages instead of only low price: UI prototype, mobile implementation, Firebase/backend integration, testing, and handover. I would use restaurant and travel apps as portfolio case studies."],
        ["Next Steps", "Complete the Freelancer profile, verify GitHub link, pin mobile app repos, prepare project descriptions, continue two bids per day, and build a simple portfolio website."],
    ]
    for row in final_rows:
        ws3.append(row)
    ws3.column_dimensions["A"].width = 28
    ws3.column_dimensions["B"].width = 110
    style_sheet(ws3)

    ws4 = wb.create_sheet("Summary")
    ws4["A1"] = "Bid Response Summary"
    ws4["A2"] = "Status"
    ws4["B2"] = "Count"
    statuses = ["Committed", "Active", "Pending", "Closed"]
    counts = {s: sum(1 for b in bids if b["status"] == s) for s in statuses}
    for s in statuses:
        ws4.append([s, counts[s]])
    ws4["D2"] = "Metric"
    ws4["E2"] = "Value"
    ws4["D3"] = "Total Bids"
    ws4["E3"] = 10
    ws4["D4"] = "Responses"
    ws4["E4"] = 5
    ws4["D5"] = "Response Rate"
    ws4["E5"] = "=E4/E3"
    ws4["E5"].number_format = "0%"
    chart = BarChart()
    chart.title = "Bid Outcomes by Status"
    chart.y_axis.title = "Count"
    chart.x_axis.title = "Status"
    data = Reference(ws4, min_col=2, min_row=2, max_row=6)
    cats = Reference(ws4, min_col=1, min_row=3, max_row=6)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.height = 7
    chart.width = 12
    ws4.add_chart(chart, "A8")
    for col in range(1, 6):
        ws4.column_dimensions[get_column_letter(col)].width = 18
    style_sheet(ws4)

    wb.save(WORKBOOK)


def build_zip():
    with ZipFile(ZIP_PATH, "w", ZIP_DEFLATED) as z:
        z.write(REPORT, REPORT.name)
        z.write(WORKBOOK, WORKBOOK.name)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_report()
    build_workbook()
    build_zip()
    print(REPORT.resolve())
    print(WORKBOOK.resolve())
    print(ZIP_PATH.resolve())


if __name__ == "__main__":
    main()
